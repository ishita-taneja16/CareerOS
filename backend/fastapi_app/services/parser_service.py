import os
import re
import json
import traceback
# pyrefly: ignore [missing-import]
import spacy
from typing import Optional, Dict, Any
# pyrefly: ignore [missing-import]
import fitz  # PyMuPDF
# pyrefly: ignore [missing-import]
from docx import Document
from models.resume_schema import ResumeSchema, ContactInfo
from config import settings
from services.llm_service import LLMService

STAGE1_PROMPT = """
You are a highly precise resume parser. Extract information from the raw resume text provided below.
Return ONLY valid JSON that matches the following schema. Keep descriptions concise. Do not add explanations.

Schema:
{{
  "contact_info": {{
    "name": "Full name",
    "email": "Email address",
    "phone": "Phone number",
    "location": "City, State or Country",
    "website": "Personal portfolio/website URL or null",
    "linkedin": "LinkedIn profile link or null",
    "github": "GitHub link or null"
  }},
  "skills": ["Skill 1", "Skill 2", ...],
  "experiences": [
    {{
      "company": "Company Name",
      "role": "Job Title",
      "location": "City, State",
      "start_date": "Start Date",
      "end_date": "End Date",
      "description_bullets": ["Bullet 1", "Bullet 2"],
      "technologies": ["Tech 1", "Tech 2"]
    }}
  ]
}}

Raw Resume Text:
---
{raw_text}
---
"""

STAGE2_PROMPT = """
You are a highly precise resume parser. Extract information from the raw resume text provided below.
Return ONLY valid JSON that matches the following schema. Keep descriptions concise. Do not add explanations.

Schema:
{{
  "education": [
    {{
      "institution": "University/School Name",
      "degree": "Degree (e.g. BS, MS)",
      "field_of_study": "Major (e.g. Computer Science)",
      "location": "City, State",
      "start_date": "Start Date",
      "end_date": "End Date",
      "gpa": "GPA or null"
    }}
  ],
  "projects": [
    {{
      "name": "Project Name",
      "description_bullets": ["Detail 1"],
      "technologies": ["Tech 1"],
      "link": "URL or null"
    }}
  ],
  "certifications": [
    {{
      "name": "Cert Name",
      "issuing_organization": "Issuer",
      "issue_date": "Date",
      "expiration_date": "Date or null"
    }}
  ],
  "achievements": ["Achievement 1", "Achievement 2"],
  "publications": ["Publication 1", "Publication 2"]
}}

Raw Resume Text:
---
{raw_text}
---
"""


class ParserService:
    def __init__(self):
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except Exception:
            # Fallback if model not downloaded
            self.nlp = None

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extracts text from PDF using PyMuPDF."""
        doc = fitz.open(file_path)
        pages_text = []
        for page in doc:
            p_text = page.get_text("text")
            if not p_text:
                p_text = page.get_text()
            pages_text.append(p_text)
        return "\n".join(pages_text)

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extracts text from DOCX using python-docx."""
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        # Read tables too
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text.append(" | ".join(row_text))
                
        return "\n".join(text)

    def extract_text(self, file_path: str) -> str:
        """Determines format and extracts raw text."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def run_llm_parser(self, raw_text: str) -> Dict[str, Any]:
        """Calls LLMService to structure the raw text into JSON using a 2-stage pipeline."""
        
        print("================================================")
        print("RAW RESUME TEXT")
        print("================================================")
        print(raw_text)

        # Stage 1: Extract contact_info, skills, experiences
        stage1_dict = {}
        stage1_failed = False
        content_stage1 = None
        
        stage1_prompt = STAGE1_PROMPT.format(raw_text=raw_text)
        try:
            content_stage1 = LLMService.call(
                prompt=stage1_prompt,
                response_format={"type": "json_object"}
            )
            print("================================================")
            print("LLM RAW RESPONSE STAGE 1")
            print("================================================")
            print(content_stage1)
            
            cleaned_stage1 = content_stage1
            if "```json" in content_stage1:
                cleaned_stage1 = content_stage1.split("```json")[1].split("```")[0].strip()
            elif "```" in content_stage1:
                cleaned_stage1 = content_stage1.split("```")[1].split("```")[0].strip()
                
            try:
                stage1_dict = json.loads(cleaned_stage1)
            except Exception as e:
                print("json.loads() failed for Stage 1. Exception:", str(e))
                print("Raw Stage 1 response:")
                print(content_stage1)
                raise e
        except Exception as e:
            print("Parser Service Stage 1 failed.")
            traceback.print_exc()
            stage1_failed = True

        # Stage 2: Extract education, projects, certifications, achievements, publications
        stage2_dict = {}
        stage2_failed = False
        content_stage2 = None
        
        stage2_prompt = STAGE2_PROMPT.format(raw_text=raw_text)
        try:
            content_stage2 = LLMService.call(
                prompt=stage2_prompt,
                response_format={"type": "json_object"}
            )
            print("================================================")
            print("LLM RAW RESPONSE STAGE 2")
            print("================================================")
            print(content_stage2)
            
            cleaned_stage2 = content_stage2
            if "```json" in content_stage2:
                cleaned_stage2 = content_stage2.split("```json")[1].split("```")[0].strip()
            elif "```" in content_stage2:
                cleaned_stage2 = content_stage2.split("```")[1].split("```")[0].strip()
                
            try:
                stage2_dict = json.loads(cleaned_stage2)
            except Exception as e:
                print("json.loads() failed for Stage 2. Exception:", str(e))
                print("Raw Stage 2 response:")
                print(content_stage2)
                raise e
        except Exception as e:
            print("Parser Service Stage 2 failed.")
            traceback.print_exc()
            stage2_failed = True

        # If both stages failed, return fallback parser dictionary
        if stage1_failed and stage2_failed:
            print("Both parsing stages failed. Returning fallback parser results.")
            return self.run_fallback_parser(raw_text)

        # Merge results
        structured_dict = {}
        
        # Helper to check if a section is populated (non-empty)
        def is_populated(val):
            if val is None:
                return False
            if isinstance(val, list) and len(val) == 0:
                return False
            if isinstance(val, dict):
                return any(is_populated(v) for v in val.values())
            if isinstance(val, str) and val.strip() == "":
                return False
            return True

        # Populate structured_dict with Stage 1 values
        for key in ["contact_info", "skills", "experiences"]:
            if key in stage1_dict:
                structured_dict[key] = stage1_dict[key]

        # Populate structured_dict with Stage 2 values
        for key in ["education", "projects", "certifications", "achievements"]:
            if key in stage2_dict:
                if key not in structured_dict or not is_populated(structured_dict[key]):
                    structured_dict[key] = stage2_dict[key]

        # Handle publications merge into achievements if ResumeSchema has no publications field
        if "publications" in stage2_dict and stage2_dict["publications"]:
            pubs = stage2_dict["publications"]
            if isinstance(pubs, list):
                if "achievements" not in structured_dict or not isinstance(structured_dict["achievements"], list):
                    structured_dict["achievements"] = []
                for pub in pubs:
                    if pub not in structured_dict["achievements"]:
                        structured_dict["achievements"].append(pub)

        # Fill missing keys with empty defaults instead of losing the entire response
        defaults = {
            "contact_info": {
                "name": "",
                "email": "",
                "phone": "",
                "location": None,
                "website": None,
                "linkedin": None,
                "github": None
            },
            "skills": [],
            "experiences": [],
            "education": [],
            "projects": [],
            "certifications": [],
            "achievements": []
        }
        
        for key, val in defaults.items():
            if key not in structured_dict or structured_dict[key] is None:
                structured_dict[key] = val
                
        return structured_dict

    def run_fallback_parser(self, raw_text: str) -> Dict[str, Any]:
        """Simple rule-based parser in case the LLM is offline or fails."""
        email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
        phone_pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        
        email_match = re.search(email_pattern, raw_text)
        phone_match = re.search(phone_pattern, raw_text)
        
        email = email_match.group(0) if email_match else ""
        phone = phone_match.group(0) if phone_match else ""
        
        name = ""
        if self.nlp:
            doc = self.nlp(raw_text[:200]) # Scan beginning for name
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    name = ent.text
                    break
        
        # If spaCy couldn't find it, take first line
        if not name:
            lines = [l.strip() for l in raw_text.split("\n") if l.strip()]
            name = lines[0] if lines else "Candidate Name"
            
        return {
            "contact_info": {
                "name": name,
                "email": email,
                "phone": phone,
                "location": "",
                "website": None,
                "linkedin": None,
                "github": None
            },
            "skills": [],
            "experiences": [],
            "education": [],
            "projects": [],
            "certifications": [],
            "achievements": []
        }

    def parse(self, file_path: str) -> ResumeSchema:
        """Orchestrates document extraction and schema loading."""
        raw_text = self.extract_text(file_path)
        structured_dict = self.run_llm_parser(raw_text)
        return ResumeSchema(**structured_dict)
