import io
# pyrefly: ignore [missing-import]
from docx import Document as DocxDocument
# pyrefly: ignore [missing-import]
from playwright.async_api import async_playwright
from models.resume_schema import ResumeSchema

class DocumentExportService:
    async def export_to_pdf(self, resume: ResumeSchema) -> bytes:
        """
        Renders structured resume data into a beautiful HTML sheet
        and prints to PDF using headless Playwright.
        """
        # Create standard, modern styling HTML template
        html_template = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    color: #333;
                    line-height: 1.4;
                    margin: 0;
                    padding: 40px;
                    font-size: 11px;
                }}
                .header {{
                    text-align: center;
                    margin-bottom: 20px;
                }}
                .name {{
                    font-size: 24px;
                    font-weight: bold;
                    margin-bottom: 5px;
                    color: #1a1a1a;
                }}
                .contact {{
                    font-size: 10px;
                    color: #666;
                    margin-bottom: 15px;
                }}
                .contact a {{
                    color: #666;
                    text-decoration: none;
                }}
                .section-title {{
                    font-size: 12px;
                    font-weight: bold;
                    text-transform: uppercase;
                    border-bottom: 1px solid #ddd;
                    margin-top: 15px;
                    margin-bottom: 8px;
                    padding-bottom: 2px;
                    color: #2c3e50;
                }}
                .skills-list {{
                    margin-bottom: 10px;
                }}
                .item-header {{
                    display: flex;
                    justify-content: space-between;
                    font-weight: bold;
                    margin-top: 8px;
                }}
                .item-sub {{
                    display: flex;
                    justify-content: space-between;
                    font-style: italic;
                    margin-bottom: 4px;
                }}
                ul {{
                    margin: 0;
                    padding-left: 20px;
                }}
                li {{
                    margin-bottom: 3px;
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="name">{resume.contact_info.name}</div>
                <div class="contact">
                    {f"Email: {resume.contact_info.email}" if resume.contact_info.email else ""}
                    {f" | Phone: {resume.contact_info.phone}" if resume.contact_info.phone else ""}
                    {f" | {resume.contact_info.location}" if resume.contact_info.location else ""}
                    <br>
                    {f"LinkedIn: <a href='{resume.contact_info.linkedin}'>{resume.contact_info.linkedin}</a>" if resume.contact_info.linkedin else ""}
                    {f" | GitHub: <a href='{resume.contact_info.github}'>{resume.contact_info.github}</a>" if resume.contact_info.github else ""}
                </div>
            </div>

            {f'<div class="section-title">Skills</div><div class="skills-list">' + ", ".join(resume.skills) + '</div>' if resume.skills else ""}

            {f'<div class="section-title">Experience</div>' if resume.experiences else ""}
            {"".join([f'''
            <div class="item-header">
                <div>{exp.company}</div>
                <div>{exp.start_date} - {exp.end_date}</div>
            </div>
            <div class="item-sub">
                <div>{exp.role}</div>
                <div>{exp.location}</div>
            </div>
            <ul>
                {"".join([f"<li>{bullet}</li>" for bullet in exp.description_bullets])}
            </ul>
            ''' for exp in resume.experiences])}

            {f'<div class="section-title">Projects</div>' if resume.projects else ""}
            {"".join([f'''
            <div class="item-header">
                <div>{proj.name}</div>
                <div>{", ".join(proj.technologies)}</div>
            </div>
            <ul>
                {"".join([f"<li>{bullet}</li>" for bullet in proj.description_bullets])}
            </ul>
            ''' for proj in resume.projects])}

            {f'<div class="section-title">Education</div>' if resume.education else ""}
            {"".join([f'''
            <div class="item-header">
                <div>{edu.institution}</div>
                <div>{edu.start_date} - {edu.end_date}</div>
            </div>
            <div class="item-sub">
                <div>{edu.degree} in {edu.field_of_study}</div>
                <div>{f"GPA: {edu.gpa}" if edu.gpa else ""}</div>
            </div>
            ''' for edu in resume.education])}

            {f'<div class="section-title">Certifications</div>' if resume.certifications else ""}
            {"".join([f'''
            <div class="item-header">
                <div>{cert.name}</div>
                <div>{cert.issue_date}</div>
            </div>
            <div class="item-sub">
                <div>{cert.issuing_organization}</div>
            </div>
            ''' for cert in resume.certifications])}

            {f'<div class="section-title">Achievements</div><ul>' + "".join([f"<li>{ach}</li>" for ach in resume.achievements]) + '</ul>' if resume.achievements else ""}
        </body>
        </html>
        """

        # Boot headless Playwright tab to print PDF
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()
            await page.set_content(html_template)
            pdf_bytes = await page.pdf(format="Letter", print_background=True)
            await browser.close()
            return pdf_bytes

    def export_to_docx(self, resume: ResumeSchema) -> bytes:
        """
        Builds a structured Word document from candidate profile details.
        """
        doc = DocxDocument()
        
        # Name and Header
        name_p = doc.add_paragraph()
        name_run = name_p.add_run(resume.contact_info.name)
        name_run.font.size = 20
        name_run.bold = True
        
        contact_p = doc.add_paragraph()
        contact_info_list = []
        if resume.contact_info.email:
            contact_info_list.append(resume.contact_info.email)
        if resume.contact_info.phone:
            contact_info_list.append(resume.contact_info.phone)
        if resume.contact_info.location:
            contact_info_list.append(resume.contact_info.location)
        contact_p.add_run(" | ".join(contact_info_list))

        # Skills
        if resume.skills:
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(resume.skills))

        # Experience
        if resume.experiences:
            doc.add_heading("Experience", level=1)
            for exp in resume.experiences:
                exp_p = doc.add_paragraph()
                exp_p.add_run(f"{exp.company} - {exp.role}").bold = True
                exp_p.add_run(f"\n{exp.start_date} - {exp.end_date} | {exp.location}").italic = True
                for bullet in exp.description_bullets:
                    doc.add_paragraph(bullet, style='List Bullet')

        # Projects
        if resume.projects:
            doc.add_heading("Projects", level=1)
            for proj in resume.projects:
                proj_p = doc.add_paragraph()
                proj_p.add_run(f"{proj.name}").bold = True
                proj_p.add_run(f"\nTechnologies: {', '.join(proj.technologies)}").italic = True
                for bullet in proj.description_bullets:
                    doc.add_paragraph(bullet, style='List Bullet')

        # Education
        if resume.education:
            doc.add_heading("Education", level=1)
            for edu in resume.education:
                edu_p = doc.add_paragraph()
                edu_p.add_run(f"{edu.institution} - {edu.degree} in {edu.field_of_study}").bold = True
                edu_p.add_run(f"\n{edu.start_date} - {edu.end_date}").italic = True

        # Save to memory stream
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return stream.getvalue()
